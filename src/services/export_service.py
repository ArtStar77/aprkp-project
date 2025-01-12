import logging
import pandas as pd
import openpyxl
from pathlib import Path
from decimal import Decimal
from datetime import datetime
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import reportlab  # Добавленный импорт
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER, legal
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, Frame, PageTemplate, BaseDocTemplate
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import json  # For loading JSON configuration
import configparser  # For loading INI configuration

from src.models.commercial_offer import CommercialOffer
from src.models.product import Product
from src.utils.formatters import Formatters

# Настройка логирования
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Настройка страницы PDF
PAGE_SIZES = {
    'A4': A4,
    'LETTER': LETTER,
    'LEGAL': legal,
    # Добавьте другие размеры при необходимости
}

class ExportService:
    @staticmethod
    def rgb_to_hex(rgb_color):
        """Convert RGB color to hex string."""
        if hasattr(rgb_color, 'rgb'):
            # Если это объект с атрибутом rgb
            rgb_value = rgb_color.rgb
            if rgb_value is None:
                return '000000'
            # Преобразуем целое число в RGB компоненты
            red = (rgb_value >> 16) & 0xFF
            green = (rgb_value >> 8) & 0xFF
            blue = rgb_value & 0xFF
            return f"{red:02x}{green:02x}{blue:02x}"
        elif isinstance(rgb_color, RGBColor):
            # Если это объект RGBColor с отдельными компонентами
            return f"{rgb_color._red:02x}{rgb_color._green:02x}{rgb_color._blue:02x}"
        else:
            # Если это кортеж или список RGB значений
            try:
                red, green, blue = rgb_color
                return f"{red:02x}{green:02x}{blue:02x}"
            except:
                return '000000'  # Возвращаем черный цвет по умолчанию

    @staticmethod
    @staticmethod
    def _set_table_format(table):
        """Apply corporate table formatting."""
        # Установка цвета фона для заголовка
        header_rgb = (40, 60, 92)  # Тёмно-синий цвет как в оригинале

        # Форматирование заголовка
        for cell in table.rows[0].cells:
            # Цвет фона
            tcPr = cell._tc.get_or_add_tcPr()
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ExportService.rgb_to_hex(header_rgb)}"/>')
            tcPr.append(shading_elm)

            # Текст
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = "Calibri Light"  # Устанавливаем шрифт Calibri Light
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # Белый текст
            run.font.size = Pt(10)

    @staticmethod
    def _add_group_header(table, name):
        """Add formatted group header row."""
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT.EXACTLY
        row.height = Pt(30)  # Высота строки для заголовка группы

        # Объединяем ячейки
        cell = row.cells[0]
        for i in range(1, len(row.cells)):
            cell.merge(row.cells[i])

        # Форматируем ячейку
        cell_rgb = (240, 240, 240)  # Светло-серый фон для заголовка группы
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ExportService.rgb_to_hex(cell_rgb)}"/>')
        tcPr.append(shading_elm)

        # Форматируем текст
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(name)
        run.font.name = "Calibri Light"  # Устанавливаем шрифт Calibri Light
        run.font.bold = True
        run.font.size = Pt(11)

    @staticmethod
    def _format_product_row(row, is_odd=False):
        """Format regular product row."""
        # Чередующийся цвет фона для строк
        if is_odd:
            bg_rgb = (248, 248, 248)  # Очень светло-серый
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ExportService.rgb_to_hex(bg_rgb)}"/>')
                tcPr.append(shading_elm)

    @staticmethod
    def _fill_table_row(table, product, item_number=None, is_odd=False):
        """Fill table row with product data and formatting."""
        if product.is_header:
            ExportService._add_group_header(table, product.name)
        else:
            row = table.add_row()
            ExportService._format_product_row(row, is_odd)

            values = [
                str(item_number) if item_number else "",
                product.name,
                str(product.quantity),
                product.unit,
                Formatters.format_currency(product.client_price, False),
                Formatters.format_currency(product.total_price, False)
            ]

            alignments = [
                WD_ALIGN_PARAGRAPH.CENTER,  # №
                WD_ALIGN_PARAGRAPH.LEFT,    # Наименование
                WD_ALIGN_PARAGRAPH.CENTER,  # Количество
                WD_ALIGN_PARAGRAPH.CENTER,  # Ед.изм.
                WD_ALIGN_PARAGRAPH.RIGHT,   # Цена
                WD_ALIGN_PARAGRAPH.RIGHT    # Сумма
            ]

            for cell, value, alignment in zip(row.cells, values, alignments):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = alignment
                run = paragraph.add_run(value)
                run.font.name = "Calibri Light"  # Устанавливаем шрифт Calibri Light
                run.font.size = Pt(10)

    @staticmethod
    def _replace_placeholders(doc, replacements):
        """Replace placeholders while preserving formatting."""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if not any(old in paragraph.text for old in replacements.keys()):
                continue

            # Store original text with replacements
            text = paragraph.text
            for old, new in replacements.items():
                if old in text:
                    text = text.replace(old, str(new))

            # Store original formatting
            if len(paragraph.runs) > 0:
                original_run = paragraph.runs[0]
                font_name = original_run.font.name
                font_size = original_run.font.size
                bold = original_run.bold
                italic = original_run.italic

                # Очищаем параграф
                for run in paragraph.runs:
                    paragraph._p.remove(run._r)

                # Добавляем новый текст с сохраненным форматированием
                new_run = paragraph.add_run(text)
                new_run.font.name = font_name
                new_run.font.size = font_size
                new_run.bold = bold
                new_run.italic = italic
            else:
                paragraph.text = text

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not any(old in paragraph.text for old in replacements.keys()):
                            continue

                        # Store original text with replacements
                        text = paragraph.text
                        for old, new in replacements.items():
                            if old in text:
                                text = text.replace(old, str(new))

                        # Store original formatting
                        if len(paragraph.runs) > 0:
                            original_run = paragraph.runs[0]
                            font_name = original_run.font.name
                            font_size = original_run.font.size
                            bold = original_run.bold
                            italic = original_run.italic

                            # Очищаем параграф
                            for run in paragraph.runs:
                                paragraph._p.remove(run._r)

                            # Добавляем новый текст с сохраненным форматированием
                            new_run = paragraph.add_run(text)
                            new_run.font.name = font_name
                            new_run.font.size = font_size
                            new_run.bold = bold
                            new_run.italic = italic
                        else:
                            paragraph.text = text

    @staticmethod
    def export_word(offer: CommercialOffer, template_path: Path, output_path: Path) -> None:
        """Export commercial offer to Word document."""
        try:
            doc = Document(template_path)

            # Prepare replacements
            items_count = len([p for p in offer.products if not p.is_header])
            total_amount = sum(p.total_price for p in offer.products if not p.is_header)
            vat_amount = (total_amount * Decimal(str(offer.vat)) / (Decimal('100') + Decimal(str(offer.vat)))).quantize(Decimal('0.01'))

            replacements = {
                '[[Номер КП]]': offer.number,
                '[[Дата КП]]': offer.date.strftime("%d.%m.%Y") if offer.date else '',
                '[[всего_позиций]]': str(items_count),
                '[[итого]]': Formatters.format_currency(total_amount, False),
                '[[ндс]]': Formatters.format_currency(vat_amount, False),
                '[[Срок поставки]]': offer.delivery_time or "3-4 недели",
                '[[Склад самовывоза]]': offer.self_pickup_warehouse or "МО, г. Люберцы, ул. Красная д 1, лит. С",
                '[[Гарантия]]': offer.warranty or "24 мес."
            }

            # Replace placeholders while preserving formatting
            ExportService._replace_placeholders(doc, replacements)

            # Find and process the main table
            for table in doc.tables:
                if len(table.rows) > 0:
                    # Clear existing rows except header
                    while len(table.rows) > 1:
                        table._element.remove(table.rows[-1]._element)

                    # Apply table formatting
                    ExportService._set_table_format(table)

                    # Add products
                    item_number = 1
                    for i, product in enumerate(offer.products):
                        ExportService._fill_table_row(
                            table,
                            product,
                            item_number if not product.is_header else None,
                            is_odd=(i % 2 == 1)  # Чередующийся цвет фона
                        )
                        if not product.is_header:
                            item_number += 1
                    break

            # Save document
            doc.save(output_path)
            logger.info(f"Word document successfully exported to {output_path}")

        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            raise

    @staticmethod
    def export_excel(offer: CommercialOffer, output_path: Path) -> None:
        """Export commercial offer to Excel file."""
        try:
            # Create workbook and worksheet
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Коммерческое предложение"

            # Стили
            header_fill = PatternFill(start_color="283C5C", end_color="283C5C", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=11)
            group_fill = PatternFill(start_color="F0F0F0", end_color="F0F0F0", fill_type="solid")
            group_font = Font(bold=True, size=11)
            alternate_fill = PatternFill(start_color="F8F8F8", end_color="F8F8F8", fill_type="solid")
            normal_font = Font(size=10)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Headers
            headers = ['№', 'Наименование продукции', 'Кол-во', 'Ед.изм.', 'Цена с НДС, руб.', 'Всего с НДС, руб.']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

            # Column widths
            column_widths = [5, 50, 10, 10, 15, 15]
            for i, width in enumerate(column_widths, 1):
                ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

            # Add products
            current_row = 2
            item_number = 1

            for i, product in enumerate(offer.products):
                if product.is_header:
                    # Merge cells for group header
                    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
                    cell = ws.cell(row=current_row, column=1)
                    cell.value = product.name
                    cell.font = group_font
                    cell.fill = group_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                else:
                    # Regular product row
                    row_data = [
                        item_number,
                        product.name,
                        product.quantity,
                        product.unit,
                        float(product.client_price),
                        float(product.total_price)
                    ]

                    for col, value in enumerate(row_data, 1):
                        cell = ws.cell(row=current_row, column=col)
                        cell.value = value
                        cell.font = normal_font
                        cell.border = border

                        # Выравнивание
                        if col in [1, 3, 4]:  # №, Кол-во, Ед.изм.
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                        elif col in [5, 6]:  # Цены
                            cell.alignment = Alignment(horizontal='right', vertical='center')
                            cell.number_format = '#,##0.00'
                        else:  # Наименование
                            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

                        # Чередующийся цвет фона
                        if i % 2 == 1:
                            cell.fill = alternate_fill

                    item_number += 1

                current_row += 1

            # Add totals
            current_row += 1
            total_amount = sum(p.total_price for p in offer.products if not p.is_header)
            vat_amount = (total_amount * Decimal(str(offer.vat)) / (Decimal('100') + Decimal(str(offer.vat)))).quantize(Decimal('0.01'))

            # Итого
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=5)
            total_cell = ws.cell(row=current_row, column=1, value="Итого:")
            total_cell.font = Font(bold=True, size=11)
            total_cell.alignment = Alignment(horizontal='right', vertical='center')

            amount_cell = ws.cell(row=current_row, column=6, value=float(total_amount))
            amount_cell.font = Font(bold=True, size=11)
            amount_cell.number_format = '#,##0.00'
            amount_cell.alignment = Alignment(horizontal='right', vertical='center')
            amount_cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # НДС
            current_row += 1
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=5)
            vat_title_cell = ws.cell(row=current_row, column=1, value="В том числе НДС:")
            vat_title_cell.font = Font(bold=True, size=11)
            vat_title_cell.alignment = Alignment(horizontal='right', vertical='center')

            vat_cell = ws.cell(row=current_row, column=6, value=float(vat_amount))
            vat_cell.font = Font(bold=True, size=11)
            vat_cell.number_format = '#,##0.00'
            vat_cell.alignment = Alignment(horizontal='right', vertical='center')
            vat_cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Add offer details
            current_row += 2
            details = [
                ("Дата КП:", offer.date.strftime("%d.%m.%Y") if offer.date else ""),
                ("Номер КП:", offer.number),
                ("Срок поставки:", offer.delivery_time or "3-4 недели"),
                ("Склад самовывоза:", offer.self_pickup_warehouse or "МО, г. Люберцы, ул. Красная д 1, лит. С"),
                ("Гарантия:", offer.warranty or "24 мес."),
                ("НДС:", f"{offer.vat}%"),
                ("Скидка от поставщика:", f"{offer.discount_from_supplier}%"),
                ("Наценка для клиента:", f"{offer.markup_for_client}%")
            ]

            for label, value in details:
                ws.cell(row=current_row, column=1, value=label).font = Font(bold=True)
                ws.cell(row=current_row, column=2, value=value)
                current_row += 1

            # Protect the worksheet
            ws.protection.sheet = True
            ws.protection.password = None  # Можно установить пароль при необходимости

            # Save workbook
            wb.save(output_path)
            logger.info(f"Excel file successfully created: {output_path}")

            pass

        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            raise

    @staticmethod
    def _create_excel_style(wb):
        """Create common styles for Excel workbook."""
        return {
            'header': {
                'font': Font(name='Arial', size=10, bold=True, color='FFFFFF'),
                'fill': PatternFill(start_color='283C5C', end_color='283C5C', fill_type='solid'),
                'alignment': Alignment(horizontal='center', vertical='center', wrap_text=True),
                'border': Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            },
            'group_header': {
                'font': Font(name='Arial', size=11, bold=True),
                'fill': PatternFill(start_color='F0F0F0', end_color='F0F0F0', fill_type='solid'),
                'alignment': Alignment(horizontal='center', vertical='center'),
                'border': Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            },
            'normal': {
                'font': Font(name='Arial', size=10),
                'alignment': Alignment(vertical='center'),
                'border': Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            },
            'currency': {
                'font': Font(name='Arial', size=10),
                'alignment': Alignment(horizontal='right', vertical='center'),
                'number_format': '#,##0.00',
                'border': Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            },
            'total': {
                'font': Font(name='Arial', size=11, bold=True),
                'alignment': Alignment(horizontal='right', vertical='center'),
                'border': Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            }
        }

    @staticmethod
    def _apply_style_to_cell(cell, style):
        """Apply style dictionary to cell."""
        cell.font = style['font']
        if 'fill' in style:
            cell.fill = style['fill']
        cell.alignment = style['alignment']
        cell.border = style['border']
        if 'number_format' in style:
            cell.number_format = style['number_format']

    @staticmethod
    def export_pdf(offer: CommercialOffer, output_path: Path, config_path: Path = None) -> None:
        """Export commercial offer to PDF file."""
        try:
            pdf_exporter = PDFExporter(config_path)
            pdf_exporter.export_pdf(offer, output_path)
            logger.info(f"PDF файл успешно создан: {output_path}")
        except Exception as e:
            logger.error(f"Ошибка при создании PDF: {str(e)}")
            raise

class PDFExporter:
    def __init__(self, config_path: Path = None):
        # Загрузка конфигурации
        self.config = self._load_config(config_path)

        # Регистрируем шрифты
        self._register_fonts()

        # Создаем стили
        self._create_styles()

    def _load_config(self, config_path: Path):
        """Load configuration from a file."""
        if config_path is None:
            # Используем настройки по умолчанию
            return self._default_config()
        else:
            if config_path.suffix == '.json':
                with open(config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            elif config_path.suffix == '.py':
                import importlib.util
                spec = importlib.util.spec_from_file_location("pdf_template_config", str(config_path))
                pdf_template_config = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(pdf_template_config)
                return pdf_template_config
            elif config_path.suffix == '.ini':
                config = configparser.ConfigParser()
                config.read(config_path, encoding='utf-8')
                return config
            else:
                logger.warning("Unknown config file format. Using default settings.")
                return self._default_config()

    def _default_config(self):
        """Default configuration settings."""
        return {
            "logo_path": "",
            "fonts": {
                "regular": "Helvetica",
                "bold": "Helvetica-Bold"
            },
            "styles": {
                "CustomNormal": {
                    "fontName": "Helvetica",
                    "fontSize": 10,
                    "leading": 12
                },
                "CustomHeading": {
                    "fontName": "Helvetica-Bold",
                    "fontSize": 12,
                    "leading": 14,
                    "alignment": 1
                }
            },
            "margins": {
                "left": 2,
                "right": 2,
                "top": 2,
                "bottom": 2
            },
            "page_size": "A4",
            "colors": {
                "header_bg": "#283C5C",
                "header_text": "#FFFFFF",
                "row_alternate_bg": "#F8F8F8"
            }
        }

    def _register_fonts(self):
        """Register required fonts for PDF generation."""
        try:
            fonts = self.config.get('fonts', {})
            pdfmetrics.registerFont(TTFont(fonts.get('regular', 'Helvetica'), 'arial.ttf'))
            pdfmetrics.registerFont(TTFont(fonts.get('bold', 'Helvetica-Bold'), 'arialbd.ttf'))
        except Exception as e:
            logger.warning(f"Не удалось загрузить шрифты Arial: {e}")
            # Используем встроенные шрифты если не удалось загрузить Arial
            self.styles = getSampleStyleSheet()

    def _create_styles(self):
        """Create styles based on configuration."""
        self.styles = getSampleStyleSheet()
        styles_config = self.config.get('styles', {})
        for style_name, style_attrs in styles_config.items():
            self.styles.add(ParagraphStyle(
                name=style_name,
                **style_attrs
            ))

    def export_pdf(self, offer: CommercialOffer, output_path: Path) -> None:
        """Export commercial offer to PDF file."""
        try:
            # Extracting configuration parameters
            margins = self.config.get('margins', {})
            left_margin = margins.get('left', 2) * cm
            right_margin = margins.get('right', 2) * cm
            top_margin = margins.get('top', 5) * cm
            bottom_margin = margins.get('bottom', 2) * cm

            # Обновленный способ получения размера страницы
            page_size_name = self.config.get('page_size', 'A4').upper()
            page_size = getattr(reportlab.lib.pagesizes, page_size_name, A4)

            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=page_size,
                rightMargin=right_margin,
                leftMargin=left_margin,
                topMargin=top_margin,
                bottomMargin=bottom_margin
            )

            elements = []

            # Add logo if exists
            logo_path = self.config.get('logo_path', '')
            if logo_path and Path(logo_path).exists():
                logo = Image(logo_path)
                logo.drawHeight = 2 * cm
                logo.drawWidth = 5 * cm
                elements.append(logo)
                elements.append(Spacer(1, 0.5 * cm))

            # Add title
            elements.append(Paragraph(
                f"КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ № {offer.number}",
                self.styles['CustomHeading']
            ))

            # Add date
            if offer.date:
                elements.append(Paragraph(
                    f"Дата: {offer.date.strftime('%d.%m.%Y')}",
                    self.styles['CustomNormal']
                ))

            elements.append(Spacer(1, 0.5 * cm))

            # Create table data
            data = [['№', 'Наименование продукции', 'Кол-во', 'Ед.изм.', 'Цена с НДС, руб.', 'Всего с НДС, руб.']]
            item_number = 1
            for i, product in enumerate(offer.products):
                if product.is_header:
                    # Group header
                    data.append([
                        '',
                        Paragraph(f"<b>{product.name}</b>", self.styles['CustomNormal']),
                        '', '', '', ''
                    ])
                else:
                    data.append([
                        str(item_number),
                        product.name,
                        str(product.quantity),
                        product.unit,
                        Formatters.format_currency(product.client_price, False),
                        Formatters.format_currency(product.total_price, False)
                    ])
                    item_number += 1

            # Create table
            col_widths = [1*cm, 8*cm, 2*cm, 2*cm, 3*cm, 3*cm]
            table = Table(data, colWidths=col_widths, repeatRows=1)

            # Table style
            colors_config = self.config.get('colors', {})
            header_bg_color = colors.HexColor(colors_config.get('header_bg', '#283C5C'))
            header_text_color = colors.HexColor(colors_config.get('header_text', '#FFFFFF'))
            row_alternate_bg_color = colors.HexColor(colors_config.get('row_alternate_bg', '#F8F8F8'))

            style = TableStyle([
                # Headers
                ('BACKGROUND', (0, 0), (-1, 0), header_bg_color),
                ('TEXTCOLOR', (0, 0), (-1, 0), header_text_color),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), self.config['fonts'].get('bold', 'Helvetica-Bold')),
                ('FONTSIZE', (0, 0), (-1, 0), 10),

                # Grid
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),

                # Alignment
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # №
                ('ALIGN', (2, 1), (3, -1), 'CENTER'),  # Количество, единицы
                ('ALIGN', (4, 1), (5, -1), 'RIGHT'),   # Цены

                # Padding
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ])

            # Add alternating row colors
            for i in range(1, len(data)):
                if i % 2 == 0:
                    style.add('BACKGROUND', (0, i), (-1, i), row_alternate_bg_color)

            table.setStyle(style)
            elements.append(table)

            # Add totals
            elements.append(Spacer(1, 0.5*cm))
            total_amount = sum(p.total_price for p in offer.products if not p.is_header)
            vat_amount = (total_amount * Decimal(str(offer.vat)) / (Decimal('100') + Decimal(str(offer.vat)))).quantize(Decimal('0.01'))

            elements.append(Paragraph(
                f"<b>Итого:</b> {Formatters.format_currency(total_amount)}",
                self.styles['CustomNormal']
            ))
            elements.append(Paragraph(
                f"<b>В том числе НДС {offer.vat}%:</b> {Formatters.format_currency(vat_amount)}",
                self.styles['CustomNormal']
            ))

            # Add offer details
            elements.append(Spacer(1, 1*cm))
            elements.append(Paragraph("<b>Условия поставки:</b>", self.styles['CustomNormal']))
            elements.append(Spacer(1, 0.3*cm))

            details = [
                f"Срок поставки: {offer.delivery_time or '3-4 недели'}",
                f"Склад самовывоза: {offer.self_pickup_warehouse or 'Адрес склада'}",
                f"Гарантия: {offer.warranty or '12 мес.'}",
                "Условия оплаты: 100% предоплата",
            ]

            for detail in details:
                elements.append(Paragraph(detail, self.styles['CustomNormal']))
                elements.append(Spacer(1, 0.2*cm))

            # Build PDF
            doc.build(elements)
            logger.info(f"PDF файл успешно создан: {output_path}")

        except Exception as e:
            logger.error(f"Ошибка при создании PDF: {str(e)}")
            raise

__all__ = ['ExportService', 'PDFExporter']
